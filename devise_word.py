from docx import Document


def remove_hash_from_word(file_path, new_file_path):
    # 加载Word文档
    doc = Document(file_path)

    # 遍历文档中的所有段落
    for paragraph in doc.paragraphs:
        # 替换段落中的'#'符号为''
        paragraph.text = paragraph.text.replace('*', '')
        paragraph.text = paragraph.text.replace('#', '')

    # 遍历文档中的所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 替换单元格中的'#'符号为''
                cell.text = cell.text.replace('#', '')

    # 保存修改后的文档
    doc.save(new_file_path)


# 使用函数
input_file_path = '1.docx'  # 输入文件的路径
output_file_path = '2.docx'  # 输出文件的路径

remove_hash_from_word(input_file_path, output_file_path)
